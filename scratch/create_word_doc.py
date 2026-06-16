from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    run = h.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.color.rgb = RGBColor(0, 51, 153)
    elif level == 2:
        run.font.color.rgb = RGBColor(0, 102, 204)

def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)

def create_doc():
    doc = Document()
    
    # Title
    title = doc.add_heading('TÀI LIỆU KỸ THUẬT: HỆ THỐNG AI GIÁM SÁT SOP', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph() # Spacing
    
    # 1. Tổng quan dự án
    add_heading(doc, '1. Tổng Quan Dự Án (AI Monitoring Hub)', 1)
    add_paragraph(doc, 'Hệ thống "AI Monitoring Hub" là một giải pháp giám sát thao tác lắp ráp của công nhân trên dây chuyền sản xuất theo thời gian thực (Real-time). Mục tiêu chính là đảm bảo công nhân tuân thủ tuyệt đối quy trình thao tác chuẩn (SOP), ngăn chặn các hành vi bỏ bước (skip step), hoặc làm sai quy trình dẫn đến sản phẩm lỗi.')
    add_paragraph(doc, 'Bài toán đặt ra: Trong môi trường nhà máy, việc giám sát thủ công hàng trăm trạm làm việc là bất khả thi. Nếu dùng các mô hình AI truyền thống (Deep Learning, LSTM) thì yêu cầu máy chủ phải có Card Đồ Họa (GPU) cực kỳ đắt đỏ, và mỗi lần thay đổi quy trình lại phải thu thập dữ liệu huấn luyện lại từ đầu.')
    add_paragraph(doc, 'Giải pháp đột phá của dự án: Chuyển đổi bài toán nhận diện hành động phức tạp thành bài toán "Giám sát Không gian" (Spatial Zone FSM). Hệ thống chỉ dùng AI (YOLO) để phát hiện bàn tay, sau đó dùng toán học và Máy trạng thái hữu hạn để kiểm tra xem tay đang tương tác ở khu vực nào, theo trình tự nào.')
    
    # 2. Kiến trúc và Công nghệ
    add_heading(doc, '2. Kiến Trúc và Công Nghệ Lõi', 1)
    add_paragraph(doc, 'Dự án được tối ưu hóa đặc biệt để chạy mượt mà trên môi trường máy chủ thông thường (CPU) mà không cần cấu hình quá cao.')
    add_bullet(doc, 'Phần cứng: Server Intel Xeon Silver 4510 (Không có GPU).')
    add_bullet(doc, 'Camera: IP Camera kết nối qua giao thức RTSP mạng LAN.')
    add_bullet(doc, 'Mô hình AI: YOLOv11 được xuất sang định dạng ONNX (ONNX Runtime CPU).')
    add_bullet(doc, 'Backend & Web: Python (Flask), SocketIO (WebSocket), MJPEG Stream.')
    add_bullet(doc, 'Database: MySQL (Connection Pool) và SQLite.')
    
    # 3. Luồng Hoạt Động (Pipeline)
    add_heading(doc, '3. Luồng Hoạt Động Chi Tiết (Pipeline)', 1)
    add_paragraph(doc, 'Hệ thống hoạt động theo một quy trình khép kín và tự động hoàn toàn, bao gồm các chặng sau:')
    
    add_heading(doc, 'Chặng 1: Thu thập và Đệm Video (Ring Buffer)', 2)
    add_paragraph(doc, 'Camera gửi luồng RTSP liên tục về server. Thay vì lưu toàn bộ video ra ổ cứng gây tràn bộ nhớ, hệ thống sử dụng cấu trúc dữ liệu đệm vòng (Ring Buffer). Hệ thống chỉ lưu khoảng 20 giây video gần nhất trên RAM. Nếu không có lỗi, video cũ sẽ bị ghi đè.')
    
    add_heading(doc, 'Chặng 2: Nhận diện bằng AI (YOLO CPU-only)', 2)
    add_paragraph(doc, 'Để tiết kiệm CPU, hệ thống không chạy AI trên mọi khung hình mà chạy cách frame (ví dụ 7.5 FPS thay vì 15 FPS). YOLO chỉ đảm nhận duy nhất việc khoanh vùng (Bounding box) bàn tay của công nhân.')
    
    add_heading(doc, 'Chặng 3: Xử lý Không Gian và Logic FSM', 2)
    add_paragraph(doc, 'Khu vực làm việc được chia thành các vùng đa giác (Polygon Zones). Tọa độ tay sẽ được đưa vào hàm hình học để kiểm tra xem đang nằm ở vùng nào. Sau đó, "Máy trạng thái" (Finite State Machine) sẽ đối chiếu thứ tự chạm vùng của công nhân với file cấu hình SOP (định dạng YAML).')
    
    add_heading(doc, 'Chặng 4: Cảnh báo và Lưu trữ', 2)
    add_paragraph(doc, 'Nếu FSM phát hiện vi phạm (ví dụ: tay vào vùng B trước khi vào vùng A), hệ thống sẽ:')
    add_bullet(doc, 'Phát âm thanh cảnh báo qua loa (Audio Alert).')
    add_bullet(doc, 'Trích xuất 20 giây video trong Ring Buffer ra thành file MP4 lưu trữ.')
    add_bullet(doc, 'Ghi thông tin lỗi vào MySQL.')
    add_bullet(doc, 'Đẩy thông báo tức thời (Real-time) lên màn hình Dashboard qua SocketIO.')
    
    # 4. Các Loại Logic Bước SOP Hỗ Trợ
    add_heading(doc, '4. Các Phép Toán Logic SOP Được Hỗ Trợ', 1)
    add_paragraph(doc, 'Hệ thống hỗ trợ cấu hình cực kỳ linh hoạt các loại thao tác của công nhân thông qua file YAML, không cần sửa mã nguồn:')
    add_bullet(doc, 'Zone Trigger: Chỉ cần đưa tay vào vùng là hoàn thành bước.')
    add_bullet(doc, 'Multi Trigger: Phải đưa tay vào/ra vùng N lần (Ví dụ: nhặt 3 con ốc).')
    add_bullet(doc, 'Stay In Zone: Phải giữ tay trong vùng tối thiểu N giây (Ví dụ: đè nắp khuôn 2 giây).')
    add_bullet(doc, 'Dual Task: Hai tay phải ở hai vùng khác nhau.')
    
    # 5. Các Tính Năng Vượt Trội (Ăn Điểm)
    add_heading(doc, '5. Các Tính Năng Vượt Trội (Dành Cho Phản Biện)', 1)
    add_paragraph(doc, 'Đây là các tính năng kỹ thuật sâu giúp hệ thống vận hành ổn định trong môi trường nhà máy thực tế, tránh báo lỗi giả (False-positive):')
    add_bullet(doc, 'Cơ chế lọc nhiễu (Dwell Filter): Ở xưởng, công nhân thao tác nhanh dễ quơ tay nhầm sang vùng khác. FSM yêu cầu tay phải nằm trọn trong vùng khoảng 0.2 - 0.5s thì mới ghi nhận, lướt qua sẽ bỏ qua.')
    add_bullet(doc, 'Thu hẹp ranh giới (Shrink Factor): Vùng nhận diện được thu nhỏ lại (ví dụ 35%) so với thực tế, bắt buộc công nhân phải vươn tay đúng vị trí trung tâm, tránh quét tay ngoài viền.')
    add_bullet(doc, 'Bám vết bằng khoảng cách Euclid: Hệ thống không dùng AI phân biệt tay trái/phải vì quá nặng. Nó dùng công thức khoảng cách Euclid (Euclidean Distance) để nối vị trí tay giữa 2 frame liên tiếp, từ đó suy ra đâu là tay trái, đâu là phải ổn định.')
    add_bullet(doc, 'Tự dọn dẹp đĩa (Storage Cleanup): Luồng nền (Daemon thread) quét ổ đĩa liên tục, nếu vượt quá 85% dung lượng sẽ tự tìm và xóa các clip vi phạm cũ, tránh gây sập server.')
    add_bullet(doc, 'Auto-Reconnect RTSP: Rớt mạng camera không làm treo hệ thống. Thread quản lý camera sẽ tự động ngủ (sleep) 5 giây và thử kết nối lại tối đa 10 lần.')
    
    # 6. Hướng Dẫn Mapping Code
    add_heading(doc, '6. Sơ Đồ Code Chi Tiết (Code Mapping)', 1)
    add_paragraph(doc, 'Khi cần tra cứu logic, hệ thống được phân bổ ở các file sau:')
    add_bullet(doc, 'Nhận luồng Video (RTSP): shared/rtsp_manager.py')
    add_bullet(doc, 'AI (YOLO ONNX): projects/sop_monitoring/hand_detector.py')
    add_bullet(doc, 'Máy trạng thái FSM lõi: projects/sop_monitoring/core/engines/TFF4040_engine.py')
    add_bullet(doc, 'Luồng điều phối (Orchestrator): projects/sop_monitoring/processor.py')
    add_bullet(doc, 'Lưu Video Ring Buffer: projects/sop_monitoring/buffer.py & shared/events/clip_saver.py')
    
    doc.save('C:/Users/it07/Downloads/AI_Monitoring_Hub/projects/sop_monitoring/docs/Bao_Cao_Chi_Tiet_SOP.docx')

if __name__ == "__main__":
    create_doc()
